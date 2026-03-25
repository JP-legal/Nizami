"""
Preview and full text extraction for PDF/DOC/DOCX and images.
- PDF/DOCX: extract full text (optionally refined via OpenAI); after extraction a summary is generated and cached; message content uses that summary.
- Images: OpenAI Vision OCR; summary is generated from OCR text and injected into message content.
- Preview: first 1-2 pages (PDF) or first sections (DOCX); images use OpenAI OCR.
"""

import base64
import logging
import os
import re

logger = logging.getLogger(__name__)

EXTRACTOR_VERSION = "2"  # Bump when extraction path changes (e.g. add LLM)
PREVIEW_PDF_PAGES = 2
PREVIEW_DOCX_PARAGRAPHS = 30  # roughly first sections


def _normalize_text(*, text: str) -> str:
    """Normalize whitespace and strip."""
    if not text:
        return ""
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_preview_from_pdf(*, file_path: str) -> str:
    """Extract text from first 1-2 pages of PDF. Use get_text(); if empty, treat as scanned (OCR not implemented here)."""
    try:
        import fitz
        doc = fitz.open(file_path)
        parts = []
        for i in range(min(PREVIEW_PDF_PAGES, len(doc))):
            page = doc[i]
            t = page.get_text()
            if t:
                parts.append(t)
        doc.close()
        return _normalize_text(text="\n".join(parts))
    except Exception as e:
        logger.warning("PDF preview extraction failed for %s: %s", file_path, e)
        return ""


def extract_full_from_pdf(*, file_path: str) -> str:
    """Extract all text from PDF."""
    try:
        import fitz
        doc = fitz.open(file_path)
        parts = []
        for page in doc:
            parts.append(page.get_text())
        doc.close()
        return _normalize_text(text="\n".join(parts))
    except Exception as e:
        logger.warning("PDF full extraction failed for %s: %s", file_path, e)
        return ""


def extract_preview_from_docx(*, file_path: str) -> str:
    """Extract first sections from DOCX."""
    try:
        import docx
        doc = docx.Document(file_path)
        paras = doc.paragraphs[:PREVIEW_DOCX_PARAGRAPHS]
        text = "\n".join(p.text for p in paras if p.text)
        return _normalize_text(text=text)
    except Exception as e:
        logger.warning("DOCX preview extraction failed for %s: %s", file_path, e)
        return ""


def extract_full_from_docx(*, file_path: str) -> str:
    """Extract all text from DOCX."""
    try:
        import docx
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return _normalize_text(text=text)
    except Exception as e:
        logger.warning("DOCX full extraction failed for %s: %s", file_path, e)
        return ""


# Chunk size for LLM extraction (chars) to stay within context
_OPENAI_EXTRACT_CHUNK_CHARS = 60_000


def _extract_text_via_openai(*, raw_text: str) -> str:
    """Send raw extracted text to OpenAI to extract/clean and return structured document text."""
    from django.conf import settings
    if not getattr(settings, "OPENAI_API_KEY", None) or not getattr(settings, "USE_OPENAI_FOR_EXTRACTION", False):
        return raw_text
    from langchain_core.messages import HumanMessage, SystemMessage
    from src.chats.utils import create_llm
    if not raw_text or not raw_text.strip():
        return raw_text
    text = raw_text.strip()
    parts = []
    for i in range(0, len(text), _OPENAI_EXTRACT_CHUNK_CHARS):
        chunk = text[i : i + _OPENAI_EXTRACT_CHUNK_CHARS]
        try:
            llm = create_llm("gpt-4o-mini", temperature=0)
            messages = [
                SystemMessage(content="You are a document text extractor. Extract and return all text from the document content. Preserve structure, paragraphs, headings, and content exactly. Return only the extracted document text, no commentary or metadata."),
                HumanMessage(content=chunk),
            ]
            response = llm.invoke(messages)
            out = (response.content or "").strip()
            if out:
                parts.append(out)
        except Exception as e:
            logger.warning("OpenAI extraction chunk failed: %s", e)
            parts.append(chunk)
    return _normalize_text(text="\n\n".join(parts)) if parts else raw_text


# Image MIME types and extensions that use OpenAI Vision OCR
_IMAGE_MIME_TYPES = frozenset({
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/gif",
    "image/webp",
})
_IMAGE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".gif", ".webp"})


def _is_image_file(*, file_path: str, mime_type: str) -> bool:
    """Return True if the file is an image we can run OCR on."""
    if mime_type and mime_type.lower() in _IMAGE_MIME_TYPES:
        return True
    if file_path:
        ext = os.path.splitext(file_path)[1].lower()
        if ext in _IMAGE_EXTENSIONS:
            return True
    return False


def _extract_image_via_openai_ocr(*, file_path: str, mime_type: str) -> str:
    """Extract text from a single image using OpenAI Vision (OCR)."""
    from django.conf import settings
    if not getattr(settings, "OPENAI_API_KEY", None) or not getattr(settings, "USE_OPENAI_FOR_EXTRACTION", False):
        return ""
    try:
        with open(file_path, "rb") as f:
            img_bytes = f.read()
    except Exception as e:
        logger.warning("Could not read image file %s: %s", file_path, e)
        return ""
    b64 = base64.standard_b64encode(img_bytes).decode("ascii")
    # Map mime or extension to data URL media type
    ext = os.path.splitext(file_path)[1].lower() if file_path else ""
    mime = (mime_type or "").strip().lower()
    if mime in ("image/jpeg", "image/jpg") or ext in (".jpg", ".jpeg"):
        media_type = "image/jpeg"
    elif mime == "image/png" or ext == ".png":
        media_type = "image/png"
    elif mime == "image/gif" or ext == ".gif":
        media_type = "image/gif"
    elif mime == "image/webp" or ext == ".webp":
        media_type = "image/webp"
    else:
        media_type = "image/png"
    content = [
        {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{b64}"}},
        {
            "type": "text",
            "text": "Extract all text from this image using OCR. Preserve order, structure, and line breaks. Return only the extracted text, no commentary.",
        },
    ]
    try:
        from langchain_core.messages import HumanMessage
        from src.chats.utils import create_llm
        llm = create_llm("gpt-4o-mini", temperature=0)
        msg = HumanMessage(content=content)
        response = llm.invoke([msg])
        return _normalize_text(text=(response.content or "").strip())
    except Exception as e:
        logger.warning("OpenAI OCR failed for image %s: %s", file_path, e)
        return ""


def _extract_pdf_via_openai_vision(*, file_path: str) -> str:
    """Render PDF pages to images and extract text via OpenAI Vision (for scanned/image PDFs)."""
    from django.conf import settings
    if not getattr(settings, "OPENAI_API_KEY", None) or not getattr(settings, "USE_OPENAI_FOR_EXTRACTION", False):
        return ""
    import fitz
    from langchain_core.messages import HumanMessage
    from src.chats.utils import create_llm
    parts = []
    try:
        doc = fitz.open(file_path)
        # Process in small batches to avoid token limits (e.g. 4 pages per request)
        pages_per_request = 4
        for start in range(0, len(doc), pages_per_request):
            content = []
            for i in range(start, min(start + pages_per_request, len(doc))):
                page = doc[i]
                pix = page.get_pixmap(dpi=150, alpha=False)
                img_bytes = pix.tobytes(output="png")
                b64 = base64.standard_b64encode(img_bytes).decode("ascii")
                content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:image/png;base64,{b64}"},
                })
            content.append({
                "type": "text",
                "text": "Extract all text from these document pages. Preserve order and structure. Return only the extracted text, one block per page if helpful.",
            })
            try:
                llm = create_llm("gpt-4o-mini", temperature=0)
                msg = HumanMessage(content=content)
                response = llm.invoke([msg])
                if response.content:
                    parts.append(response.content.strip())
            except Exception as e:
                logger.warning("OpenAI vision extraction batch failed: %s", e)
        doc.close()
    except Exception as e:
        logger.warning("PDF vision extraction failed for %s: %s", file_path, e)
    return _normalize_text(text="\n\n".join(parts)) if parts else ""


def get_preview_text(*, file_path: str, mime_type: str) -> str:
    """Return preview text based on file type. Images use OpenAI OCR."""
    if mime_type == "application/pdf" or (file_path or "").lower().endswith(".pdf"):
        return extract_preview_from_pdf(file_path=file_path)
    if (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or (file_path or "").lower().endswith(".docx")
    ):
        return extract_preview_from_docx(file_path=file_path)
    if mime_type == "application/msword" or (file_path or "").lower().endswith(".doc"):
        return extract_preview_from_docx(file_path=file_path)
    if _is_image_file(file_path=file_path, mime_type=mime_type or ""):
        return _extract_image_via_openai_ocr(file_path=file_path, mime_type=mime_type or "")
    return ""


def get_full_text(*, file_path: str, mime_type: str) -> str:
    """Return full extracted text. Uses library extraction first; when USE_OPENAI_FOR_EXTRACTION is True, refines via OpenAI LLM. Images use OpenAI Vision OCR."""
    from django.conf import settings
    use_openai = getattr(settings, "USE_OPENAI_FOR_EXTRACTION", False)
    from_vision = False
    mime = mime_type or ""

    raw = ""
    if _is_image_file(file_path=file_path, mime_type=mime):
        raw = _extract_image_via_openai_ocr(file_path=file_path, mime_type=mime)
        from_vision = bool(raw)
    elif mime_type == "application/pdf" or (file_path or "").lower().endswith(".pdf"):
        raw = extract_full_from_pdf(file_path=file_path)
        if use_openai and not raw.strip():
            raw = _extract_pdf_via_openai_vision(file_path=file_path)
            from_vision = bool(raw)
    elif (
        mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or (file_path or "").lower().endswith(".docx")
    ):
        raw = extract_full_from_docx(file_path=file_path)
    elif mime_type == "application/msword" or (file_path or "").lower().endswith(".doc"):
        try:
            raw = extract_full_from_docx(file_path=file_path)
        except Exception:
            pass
        if not raw:
            try:
                from src.common.text_extraction import extract_text_from_file
                raw = extract_text_from_file(file_path=file_path)
            except Exception as e:
                logger.warning("DOC full extraction failed for %s: %s", file_path, e)
        raw = _normalize_text(text=raw) if raw else ""

    if use_openai and raw and not from_vision:
        return _extract_text_via_openai(raw_text=raw)
    return raw or ""
